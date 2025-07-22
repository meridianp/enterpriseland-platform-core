"""Template service for document template management."""

import os
import tempfile
import zipfile
from typing import Dict, Any, List, Optional
from string import Template
from jinja2 import Template as JinjaTemplate, Environment, FileSystemLoader
from docx import Document as DocxDocument
from docx.shared import Pt
from openpyxl import load_workbook
from django.conf import settings
from django.template import Context, Template as DjangoTemplate

from ..models import Document, DocumentTemplate, TemplateField, Folder
from .storage_service import StorageService
from .document_service import DocumentService


class TemplateService:
    """Service for managing document templates."""
    
    def __init__(self):
        self.storage_service = StorageService()
        self.document_service = DocumentService()
        
        # Jinja2 environment for advanced templating
        self.jinja_env = Environment(
            loader=FileSystemLoader(
                getattr(settings, 'DOCUMENTS_TEMPLATE_DIR', '/tmp')
            ),
            autoescape=True
        )
    
    def create_document_from_template(
        self,
        template: DocumentTemplate,
        user,
        data: Dict[str, Any],
        folder: Optional[Folder] = None,
        name: Optional[str] = None
    ) -> Document:
        """Create a new document from a template with data substitution."""
        # Validate data against template fields
        self._validate_template_data(template, data)
        
        # Download template file
        template_content = self.storage_service.download_file(template.file_path)
        
        # Process based on template type
        if template.file_type == 'docx':
            processed_content = self._process_docx_template(template_content, data)
            file_extension = 'docx'
        elif template.file_type == 'xlsx':
            processed_content = self._process_xlsx_template(template_content, data)
            file_extension = 'xlsx'
        elif template.file_type == 'html':
            processed_content = self._process_html_template(template_content, data)
            file_extension = 'html'
        elif template.file_type == 'txt':
            processed_content = self._process_text_template(template_content, data)
            file_extension = 'txt'
        else:
            raise ValueError(f"Unsupported template type: {template.file_type}")
        
        # Generate document name
        if not name:
            name = self._generate_document_name(template, data)
        
        # Create document
        with tempfile.NamedTemporaryFile(suffix=f'.{file_extension}', delete=False) as tmp_file:
            tmp_file.write(processed_content)
            tmp_file.flush()
            
            # Create document using document service
            with open(tmp_file.name, 'rb') as f:
                document = self.document_service.create_document(
                    file=f,
                    name=name,
                    user=user,
                    folder=folder,
                    description=f"Created from template: {template.name}",
                    category=template.category,
                    tags=['template-generated', template.name]
                )
            
            # Clean up
            os.unlink(tmp_file.name)
        
        # Update template usage
        template.increment_usage()
        
        return document
    
    def _validate_template_data(self, template: DocumentTemplate, data: Dict[str, Any]) -> None:
        """Validate data against template field requirements."""
        errors = []
        
        for field in template.fields.filter(is_active=True):
            field_name = field.name
            field_value = data.get(field_name)
            
            # Check required fields
            if field.is_required and not field_value:
                errors.append(f"Field '{field.label}' is required")
                continue
            
            # Validate field value if provided
            if field_value:
                try:
                    field.validate_value(field_value)
                except ValueError as e:
                    errors.append(str(e))
        
        if errors:
            raise ValueError(f"Template validation errors: {'; '.join(errors)}")
    
    def _process_docx_template(self, template_content: bytes, data: Dict[str, Any]) -> bytes:
        """Process Word document template."""
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp_template:
            tmp_template.write(template_content)
            tmp_template.flush()
            
            # Open document
            doc = DocxDocument(tmp_template.name)
            
            # Replace placeholders in paragraphs
            for paragraph in doc.paragraphs:
                self._replace_paragraph_text(paragraph, data)
            
            # Replace placeholders in tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            self._replace_paragraph_text(paragraph, data)
            
            # Replace in headers and footers
            for section in doc.sections:
                # Header
                for paragraph in section.header.paragraphs:
                    self._replace_paragraph_text(paragraph, data)
                
                # Footer
                for paragraph in section.footer.paragraphs:
                    self._replace_paragraph_text(paragraph, data)
            
            # Save processed document
            with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp_output:
                doc.save(tmp_output.name)
                
                # Read processed content
                with open(tmp_output.name, 'rb') as f:
                    processed_content = f.read()
                
                # Clean up
                os.unlink(tmp_output.name)
            
            # Clean up template
            os.unlink(tmp_template.name)
            
            return processed_content
    
    def _replace_paragraph_text(self, paragraph, data: Dict[str, Any]) -> None:
        """Replace placeholders in a paragraph while preserving formatting."""
        if '{{' in paragraph.text:
            # Get full text
            full_text = paragraph.text
            
            # Replace placeholders
            for key, value in data.items():
                placeholder = f"{{{{{key}}}}}"
                if placeholder in full_text:
                    full_text = full_text.replace(placeholder, str(value))
            
            # Clear paragraph and add new text
            paragraph.clear()
            paragraph.add_run(full_text)
    
    def _process_xlsx_template(self, template_content: bytes, data: Dict[str, Any]) -> bytes:
        """Process Excel template."""
        with tempfile.NamedTemporaryFile(suffix='.xlsx', delete=False) as tmp_template:
            tmp_template.write(template_content)
            tmp_template.flush()
            
            # Open workbook
            workbook = load_workbook(tmp_template.name)
            
            # Process all sheets
            for sheet in workbook.worksheets:
                # Iterate through all cells
                for row in sheet.iter_rows():
                    for cell in row:
                        if cell.value and isinstance(cell.value, str) and '{{' in cell.value:
                            # Replace placeholders
                            new_value = cell.value
                            for key, value in data.items():
                                placeholder = f"{{{{{key}}}}}"
                                if placeholder in new_value:
                                    new_value = new_value.replace(placeholder, str(value))
                            
                            # Update cell value
                            cell.value = new_value
            
            # Save processed workbook
            with tempfile.NamedTemporaryFile(suffix='.xlsx', delete=False) as tmp_output:
                workbook.save(tmp_output.name)
                
                # Read processed content
                with open(tmp_output.name, 'rb') as f:
                    processed_content = f.read()
                
                # Clean up
                os.unlink(tmp_output.name)
            
            # Clean up template
            os.unlink(tmp_template.name)
            
            return processed_content
    
    def _process_html_template(self, template_content: bytes, data: Dict[str, Any]) -> bytes:
        """Process HTML template using Jinja2."""
        try:
            # Decode template
            template_str = template_content.decode('utf-8')
            
            # Create Jinja2 template
            template = JinjaTemplate(template_str)
            
            # Render with data
            rendered = template.render(**data)
            
            return rendered.encode('utf-8')
        
        except Exception as e:
            raise Exception(f"Error processing HTML template: {str(e)}")
    
    def _process_text_template(self, template_content: bytes, data: Dict[str, Any]) -> bytes:
        """Process plain text template."""
        try:
            # Decode template
            template_str = template_content.decode('utf-8')
            
            # Use Python string template for simple substitution
            template = Template(template_str)
            
            # Convert data keys to use $ prefix
            template_data = {k: str(v) for k, v in data.items()}
            
            # Render
            rendered = template.safe_substitute(**template_data)
            
            return rendered.encode('utf-8')
        
        except Exception as e:
            raise Exception(f"Error processing text template: {str(e)}")
    
    def _generate_document_name(self, template: DocumentTemplate, data: Dict[str, Any]) -> str:
        """Generate document name based on template and data."""
        # Try to use template naming pattern
        if hasattr(template, 'naming_pattern') and template.naming_pattern:
            try:
                name_template = Template(template.naming_pattern)
                name = name_template.safe_substitute(**data)
                return name
            except:
                pass
        
        # Default naming
        import datetime
        timestamp = datetime.datetime.now().strftime('%Y%m%d_%H%M%S')
        
        # Try to use a meaningful field
        for field in ['name', 'title', 'subject', 'reference']:
            if field in data and data[field]:
                return f"{data[field]}_{timestamp}"
        
        # Fallback to template name
        return f"{template.name}_{timestamp}"
    
    def create_template_from_document(
        self,
        document: Document,
        name: str,
        description: str,
        category: str,
        fields: List[Dict[str, Any]],
        user
    ) -> DocumentTemplate:
        """Create a template from an existing document."""
        # Download document
        document_content = self.storage_service.download_file(document.file_path)
        
        # Analyze document for placeholders
        placeholders = self._extract_placeholders(document_content, document.file_extension)
        
        # Create template
        template_path = self.storage_service.generate_path(
            file_name=f"template_{document.file_name}",
            user=user
        )
        
        # Upload template file
        upload_result = self.storage_service.upload_file(
            document_content,
            template_path,
            document.mime_type
        )
        
        # Create template record
        template = DocumentTemplate.objects.create(
            name=name,
            description=description,
            category=category,
            file_path=upload_result['path'],
            file_type=document.file_extension,
            group=user.group,
            created_by=user,
            modified_by=user
        )
        
        # Create fields
        for field_data in fields:
            TemplateField.objects.create(
                template=template,
                **field_data
            )
        
        # Auto-create fields from placeholders
        for placeholder in placeholders:
            if not template.fields.filter(name=placeholder).exists():
                TemplateField.objects.create(
                    template=template,
                    name=placeholder,
                    label=placeholder.replace('_', ' ').title(),
                    field_type='text',
                    template_variable=f"{{{{{placeholder}}}}}"
                )
        
        return template
    
    def _extract_placeholders(self, content: bytes, file_type: str) -> List[str]:
        """Extract placeholder names from template content."""
        placeholders = set()
        
        try:
            if file_type in ['txt', 'html']:
                # Simple text search
                text = content.decode('utf-8', errors='ignore')
                import re
                
                # Find {{placeholder}} pattern
                pattern = r'\{\{(\w+)\}\}'
                matches = re.findall(pattern, text)
                placeholders.update(matches)
                
                # Find ${placeholder} pattern
                pattern = r'\$\{(\w+)\}'
                matches = re.findall(pattern, text)
                placeholders.update(matches)
            
            elif file_type == 'docx':
                # Extract from Word document
                with tempfile.NamedTemporaryFile(suffix='.docx') as tmp:
                    tmp.write(content)
                    tmp.flush()
                    
                    doc = DocxDocument(tmp.name)
                    
                    # Check paragraphs
                    for paragraph in doc.paragraphs:
                        if '{{' in paragraph.text:
                            import re
                            pattern = r'\{\{(\w+)\}\}'
                            matches = re.findall(pattern, paragraph.text)
                            placeholders.update(matches)
        
        except Exception:
            pass
        
        return list(placeholders)
    
    def preview_template(
        self,
        template: DocumentTemplate,
        sample_data: Optional[Dict[str, Any]] = None
    ) -> Dict[str, Any]:
        """Preview template with sample data."""
        if not sample_data:
            # Generate sample data from fields
            sample_data = {}
            for field in template.fields.filter(is_active=True):
                if field.default_value:
                    sample_data[field.name] = field.default_value
                else:
                    # Generate sample based on type
                    if field.field_type == 'text':
                        sample_data[field.name] = f"Sample {field.label}"
                    elif field.field_type == 'number':
                        sample_data[field.name] = "123"
                    elif field.field_type == 'date':
                        sample_data[field.name] = "2024-01-01"
                    elif field.field_type == 'email':
                        sample_data[field.name] = "sample@example.com"
                    elif field.field_type == 'boolean':
                        sample_data[field.name] = "Yes"
        
        # Process template with sample data
        try:
            # Download template
            template_content = self.storage_service.download_file(template.file_path)
            
            # Process based on type
            if template.file_type == 'html':
                processed = self._process_html_template(template_content, sample_data)
                preview_html = processed.decode('utf-8')
            else:
                # For other types, show the filled data
                preview_html = "<h3>Preview Data:</h3><dl>"
                for key, value in sample_data.items():
                    preview_html += f"<dt><strong>{key}:</strong></dt><dd>{value}</dd>"
                preview_html += "</dl>"
            
            return {
                'success': True,
                'preview_html': preview_html,
                'sample_data': sample_data
            }
        
        except Exception as e:
            return {
                'success': False,
                'error': str(e)
            }