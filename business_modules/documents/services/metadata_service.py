"""Metadata extraction service for documents."""

import os
import tempfile
from datetime import datetime
from typing import Dict, Any, Optional
from PIL import Image
from PIL.ExifTags import TAGS, GPSTAGS
import PyPDF2
import openpyxl
from docx import Document as DocxDocument
import magic
import mutagen
from django.conf import settings

from ..models import Document, DocumentMetadata
from .storage_service import StorageService


class MetadataService:
    """Service for extracting and managing document metadata."""
    
    def __init__(self):
        self.storage_service = StorageService()
        self.mime = magic.Magic(mime=True)
    
    def extract_metadata(self, document: Document) -> Dict[str, Any]:
        """Extract metadata from document based on file type."""
        try:
            # Download file content
            file_content = self.storage_service.download_file(document.file_path)
            
            # Create temporary file
            with tempfile.NamedTemporaryFile(suffix=f'.{document.file_extension}', delete=False) as tmp_file:
                tmp_file.write(file_content)
                tmp_file.flush()
                tmp_path = tmp_file.name
            
            try:
                # Extract based on file type
                if document.file_extension == 'pdf':
                    metadata = self._extract_pdf_metadata(tmp_path)
                elif document.file_extension in ['jpg', 'jpeg', 'png', 'tiff', 'bmp']:
                    metadata = self._extract_image_metadata(tmp_path)
                elif document.file_extension in ['doc', 'docx']:
                    metadata = self._extract_word_metadata(tmp_path)
                elif document.file_extension in ['xls', 'xlsx']:
                    metadata = self._extract_excel_metadata(tmp_path)
                elif document.file_extension in ['mp3', 'mp4', 'avi', 'mov']:
                    metadata = self._extract_media_metadata(tmp_path)
                else:
                    metadata = self._extract_basic_metadata(tmp_path)
                
                # Add common metadata
                metadata['file_size'] = document.size
                metadata['mime_type'] = document.mime_type
                metadata['file_extension'] = document.file_extension
                
                # Update document metadata
                if hasattr(document, 'metadata'):
                    document.metadata.update_from_extraction(metadata)
                else:
                    DocumentMetadata.objects.create(
                        document=document,
                        **self._map_to_model_fields(metadata)
                    )
                
                return metadata
            
            finally:
                # Clean up
                os.unlink(tmp_path)
        
        except Exception as e:
            print(f"Error extracting metadata for document {document.id}: {e}")
            return {'error': str(e)}
    
    def _extract_pdf_metadata(self, file_path: str) -> Dict[str, Any]:
        """Extract metadata from PDF files."""
        metadata = {'raw_metadata': {}}
        
        try:
            with open(file_path, 'rb') as pdf_file:
                pdf_reader = PyPDF2.PdfReader(pdf_file)
                
                # Get document info
                if pdf_reader.metadata:
                    info = pdf_reader.metadata
                    
                    metadata['title'] = self._safe_get(info, '/Title')
                    metadata['author'] = self._safe_get(info, '/Author')
                    metadata['subject'] = self._safe_get(info, '/Subject')
                    metadata['keywords'] = self._safe_get(info, '/Keywords')
                    metadata['creator_tool'] = self._safe_get(info, '/Creator')
                    metadata['producer'] = self._safe_get(info, '/Producer')
                    
                    # Parse dates
                    creation_date = self._safe_get(info, '/CreationDate')
                    if creation_date:
                        metadata['creation_date'] = self._parse_pdf_date(creation_date)
                    
                    mod_date = self._safe_get(info, '/ModDate')
                    if mod_date:
                        metadata['modification_date'] = self._parse_pdf_date(mod_date)
                    
                    # Store raw metadata
                    metadata['raw_metadata'] = {k: str(v) for k, v in info.items()}
                
                # Get page count
                metadata['page_count'] = len(pdf_reader.pages)
                
                # Extract text for word count (first few pages)
                text = ""
                for i in range(min(5, len(pdf_reader.pages))):
                    page_text = pdf_reader.pages[i].extract_text()
                    text += page_text + " "
                
                if text.strip():
                    words = text.split()
                    metadata['word_count'] = len(words)
                    metadata['character_count'] = len(text)
        
        except Exception as e:
            metadata['error'] = str(e)
        
        return metadata
    
    def _extract_image_metadata(self, file_path: str) -> Dict[str, Any]:
        """Extract metadata from image files."""
        metadata = {'raw_metadata': {}}
        
        try:
            with Image.open(file_path) as img:
                # Basic image properties
                metadata['width'] = img.width
                metadata['height'] = img.height
                metadata['format'] = img.format
                metadata['mode'] = img.mode
                
                # EXIF data
                exif_data = img.getexif()
                if exif_data:
                    exif_dict = {}
                    
                    for tag_id, value in exif_data.items():
                        tag = TAGS.get(tag_id, tag_id)
                        exif_dict[tag] = value
                        
                        # Extract specific fields
                        if tag == 'DateTime':
                            metadata['creation_date'] = self._parse_exif_date(value)
                        elif tag == 'DateTimeOriginal':
                            metadata['creation_date'] = self._parse_exif_date(value)
                        elif tag == 'Artist':
                            metadata['author'] = value
                        elif tag == 'ImageDescription':
                            metadata['title'] = value
                        elif tag == 'Software':
                            metadata['creator_tool'] = value
                        elif tag == 'Make':
                            metadata['camera_make'] = value
                        elif tag == 'Model':
                            metadata['camera_model'] = value
                    
                    metadata['raw_metadata']['exif'] = exif_dict
                    
                    # GPS data
                    gps_info = exif_data.get(34853)  # GPSInfo tag
                    if gps_info:
                        gps_data = {}
                        for key, value in gps_info.items():
                            tag = GPSTAGS.get(key, key)
                            gps_data[tag] = value
                        
                        # Extract coordinates
                        lat = self._get_gps_coordinate(gps_data, 'GPSLatitude', 'GPSLatitudeRef')
                        lon = self._get_gps_coordinate(gps_data, 'GPSLongitude', 'GPSLongitudeRef')
                        
                        if lat is not None and lon is not None:
                            metadata['latitude'] = lat
                            metadata['longitude'] = lon
                        
                        metadata['raw_metadata']['gps'] = gps_data
        
        except Exception as e:
            metadata['error'] = str(e)
        
        return metadata
    
    def _extract_word_metadata(self, file_path: str) -> Dict[str, Any]:
        """Extract metadata from Word documents."""
        metadata = {'raw_metadata': {}}
        
        try:
            doc = DocxDocument(file_path)
            
            # Core properties
            core_props = doc.core_properties
            
            metadata['title'] = core_props.title or ''
            metadata['author'] = core_props.author or ''
            metadata['subject'] = core_props.subject or ''
            metadata['keywords'] = core_props.keywords or ''
            metadata['comments'] = core_props.comments or ''
            
            if core_props.created:
                metadata['creation_date'] = core_props.created
            if core_props.modified:
                metadata['modification_date'] = core_props.modified
            
            metadata['revision'] = core_props.revision
            metadata['last_modified_by'] = core_props.last_modified_by or ''
            
            # Document statistics
            word_count = 0
            char_count = 0
            
            for paragraph in doc.paragraphs:
                text = paragraph.text
                words = text.split()
                word_count += len(words)
                char_count += len(text)
            
            metadata['word_count'] = word_count
            metadata['character_count'] = char_count
            metadata['paragraph_count'] = len(doc.paragraphs)
            
            # Store raw properties
            metadata['raw_metadata'] = {
                'title': metadata['title'],
                'author': metadata['author'],
                'subject': metadata['subject'],
                'keywords': metadata['keywords']
            }
        
        except Exception as e:
            metadata['error'] = str(e)
        
        return metadata
    
    def _extract_excel_metadata(self, file_path: str) -> Dict[str, Any]:
        """Extract metadata from Excel files."""
        metadata = {'raw_metadata': {}}
        
        try:
            workbook = openpyxl.load_workbook(file_path, read_only=True, data_only=True)
            
            # Properties
            props = workbook.properties
            
            metadata['title'] = props.title or ''
            metadata['author'] = props.creator or ''
            metadata['subject'] = props.subject or ''
            metadata['keywords'] = props.keywords or ''
            metadata['description'] = props.description or ''
            metadata['last_modified_by'] = props.lastModifiedBy or ''
            
            if props.created:
                metadata['creation_date'] = props.created
            if props.modified:
                metadata['modification_date'] = props.modified
            
            # Workbook statistics
            metadata['sheet_count'] = len(workbook.sheetnames)
            metadata['sheet_names'] = workbook.sheetnames
            
            # Cell count (approximate)
            total_cells = 0
            for sheet in workbook:
                total_cells += sheet.max_row * sheet.max_column
            
            metadata['cell_count'] = total_cells
            
            # Store raw properties
            metadata['raw_metadata'] = {
                'title': metadata['title'],
                'sheets': metadata['sheet_names']
            }
            
            workbook.close()
        
        except Exception as e:
            metadata['error'] = str(e)
        
        return metadata
    
    def _extract_media_metadata(self, file_path: str) -> Dict[str, Any]:
        """Extract metadata from media files."""
        metadata = {'raw_metadata': {}}
        
        try:
            # Use mutagen for audio/video metadata
            file_info = mutagen.File(file_path)
            
            if file_info:
                # Duration
                if hasattr(file_info.info, 'length'):
                    metadata['duration'] = int(file_info.info.length)
                
                # Bitrate
                if hasattr(file_info.info, 'bitrate'):
                    metadata['bitrate'] = file_info.info.bitrate
                
                # Common tags
                if file_info.tags:
                    tags = {}
                    
                    # Try to get common tags
                    title = file_info.tags.get('TIT2') or file_info.tags.get('title')
                    if title:
                        metadata['title'] = str(title[0]) if isinstance(title, list) else str(title)
                    
                    artist = file_info.tags.get('TPE1') or file_info.tags.get('artist')
                    if artist:
                        metadata['author'] = str(artist[0]) if isinstance(artist, list) else str(artist)
                    
                    album = file_info.tags.get('TALB') or file_info.tags.get('album')
                    if album:
                        metadata['album'] = str(album[0]) if isinstance(album, list) else str(album)
                    
                    date = file_info.tags.get('TDRC') or file_info.tags.get('date')
                    if date:
                        try:
                            date_str = str(date[0]) if isinstance(date, list) else str(date)
                            metadata['creation_date'] = datetime.strptime(date_str[:4], '%Y')
                        except:
                            pass
                    
                    # Store all tags
                    for key, value in file_info.tags.items():
                        tags[key] = str(value)
                    
                    metadata['raw_metadata']['tags'] = tags
                
                # Format info
                if hasattr(file_info.info, 'codec'):
                    metadata['codec'] = file_info.info.codec
                
                if hasattr(file_info.info, 'channels'):
                    metadata['channels'] = file_info.info.channels
                
                if hasattr(file_info.info, 'sample_rate'):
                    metadata['sample_rate'] = file_info.info.sample_rate
        
        except Exception as e:
            metadata['error'] = str(e)
        
        return metadata
    
    def _extract_basic_metadata(self, file_path: str) -> Dict[str, Any]:
        """Extract basic metadata for unsupported file types."""
        metadata = {}
        
        try:
            import os
            stat = os.stat(file_path)
            
            metadata['file_size'] = stat.st_size
            metadata['creation_date'] = datetime.fromtimestamp(stat.st_ctime)
            metadata['modification_date'] = datetime.fromtimestamp(stat.st_mtime)
            
            # Try to get MIME type
            metadata['detected_mime'] = self.mime.from_file(file_path)
        
        except Exception as e:
            metadata['error'] = str(e)
        
        return metadata
    
    def _safe_get(self, dict_obj: Dict, key: str, default: Any = None) -> Any:
        """Safely get value from dictionary."""
        try:
            value = dict_obj.get(key, default)
            if value and hasattr(value, 'decode'):
                return value.decode('utf-8', errors='ignore')
            return value
        except:
            return default
    
    def _parse_pdf_date(self, date_string: str) -> Optional[datetime]:
        """Parse PDF date format (D:YYYYMMDDHHmmSS)."""
        try:
            if date_string.startswith('D:'):
                date_string = date_string[2:]
            
            # Extract basic date components
            year = int(date_string[0:4])
            month = int(date_string[4:6])
            day = int(date_string[6:8])
            
            # Try to get time components
            hour = int(date_string[8:10]) if len(date_string) >= 10 else 0
            minute = int(date_string[10:12]) if len(date_string) >= 12 else 0
            second = int(date_string[12:14]) if len(date_string) >= 14 else 0
            
            return datetime(year, month, day, hour, minute, second)
        except:
            return None
    
    def _parse_exif_date(self, date_string: str) -> Optional[datetime]:
        """Parse EXIF date format (YYYY:MM:DD HH:MM:SS)."""
        try:
            return datetime.strptime(date_string, '%Y:%m:%d %H:%M:%S')
        except:
            return None
    
    def _get_gps_coordinate(self, gps_data: Dict, coord_key: str, ref_key: str) -> Optional[float]:
        """Extract GPS coordinate from EXIF data."""
        try:
            coord = gps_data.get(coord_key)
            ref = gps_data.get(ref_key)
            
            if coord and ref:
                # Convert to decimal degrees
                degrees = coord[0]
                minutes = coord[1]
                seconds = coord[2]
                
                decimal = degrees + (minutes / 60.0) + (seconds / 3600.0)
                
                # Apply reference (N/S for latitude, E/W for longitude)
                if ref in ['S', 'W']:
                    decimal = -decimal
                
                return round(decimal, 8)
        except:
            return None
    
    def _map_to_model_fields(self, metadata: Dict[str, Any]) -> Dict[str, Any]:
        """Map extracted metadata to model fields."""
        return {
            'file_metadata': metadata.get('raw_metadata', {}),
            'title': metadata.get('title', '')[:500],
            'author': metadata.get('author', '')[:255],
            'subject': metadata.get('subject', '')[:500],
            'keywords': metadata.get('keywords', ''),
            'creation_date': metadata.get('creation_date'),
            'modification_date': metadata.get('modification_date'),
            'page_count': metadata.get('page_count'),
            'word_count': metadata.get('word_count'),
            'character_count': metadata.get('character_count'),
            'width': metadata.get('width'),
            'height': metadata.get('height'),
            'duration': metadata.get('duration'),
            'latitude': metadata.get('latitude'),
            'longitude': metadata.get('longitude'),
            'producer': metadata.get('producer', '')[:255],
            'creator_tool': metadata.get('creator_tool', '')[:255],
            'extraction_status': 'completed',
            'extracted_at': datetime.now()
        }